from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Import for text alignment
from bs4 import BeautifulSoup
import pdfkit
import re
import os
import markdown

def add_line_breaks_between_list_items(md_text):
    md_text = re.sub(r'(\n\d+\. .+)(\n\d+\. .+)', r'\1\n\2', md_text)
    md_text = re.sub(r'(\n  - .+)(\n  - .+)', r'\1\n\2', md_text)
    return md_text

def convert_markdown_to_pdf_and_word(markdown_file_path, pdf_file_path, word_file_path):
    os.makedirs(os.path.dirname(pdf_file_path), exist_ok=True)
    os.makedirs(os.path.dirname(word_file_path), exist_ok=True)

    with open(markdown_file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    html_text = markdown.markdown(md_text)

    doc = Document()

    soup = BeautifulSoup(html_text, 'html.parser')

    for element in soup:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            p = doc.add_paragraph()
            run = p.add_run(element.text)
            run.bold = True
            run.font.size = Pt(16 if element.name in ['h1', 'h2'] else 14)
        elif element.name == 'p':
            p = doc.add_paragraph(element.text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Set text alignment to justify
        elif element.name == 'ul':
            for item in element.find_all('li'):
                p = doc.add_paragraph()
                p.add_run('• ')
                p.add_run(item.text)
        elif element.name == 'ol':
            for idx, item in enumerate(element.find_all('li'), 1):
                p = doc.add_paragraph()
                p.add_run(f'{idx}. {item.text}')

    doc.save(word_file_path)

    markdown_text = add_line_breaks_between_list_items(md_text)
    html_text = markdown.markdown(markdown_text)

    wrapped_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: Arial, sans-serif;
                text-align: justify;  # Set text alignment to justify
            }}
            ol, ul {{
                margin-bottom: 20px;
            }}
            li {{
                margin-bottom: 5px;
            }}
        </style>
    </head>
    <body>
        {html_text}
    </body>
    </html>
    """

    pdfkit.from_string(wrapped_html, pdf_file_path, {
        'margin-top': '20mm',
        'margin-bottom': '15mm',
        'margin-right': '10mm',
        'margin-left': '10mm',
    })