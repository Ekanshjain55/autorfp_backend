from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import pdfkit
import re
import os
import markdown
import base64
from docx.enum.text import WD_BREAK


def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode()


def add_line_breaks_between_list_items(md_text):
    md_text = re.sub(r'(\n\d+\. .+)(\n\d+\. .+)', r'\1\n\2', md_text)
    md_text = re.sub(r'(\n  - .+)(\n  - .+)', r'\1\n\2', md_text)
    return md_text

def extract_project_name_from_markdown(md_text):
    # Using regex to find the pattern of the project name in Markdown
    match = re.search(r'^### PROJECT NAME\s*\n(.+)$', md_text, re.MULTILINE)
    if match:
        return match.group(1).strip()
    return None

def add_title_page_to_word(doc, project_name, logo_path):
    # logo centrally aligned
    logo_paragraph = doc.add_paragraph()
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.5))
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # title centrally aligned
    title = doc.add_paragraph()
    title_run = title.add_run(project_name)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # versioning table
    table = doc.add_table(rows=2, cols=2)  
    table.autofit = True

    # Header cells
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Version'
    hdr_cells[1].text = 'Date'

    # First row cells
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '1.0'
    row1_cells[1].text = 'dd/mm/yyyy'

    # New page
    section_break = doc.add_paragraph()
    section_break.add_run()
    section_break.add_run().add_break(WD_BREAK.PAGE)  # Adding page break


def add_footer_to_word(doc, company_name):
    # company name to the footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = company_name

def add_title_page_to_pdf(html_text, project_name, logo_path):
    logo_base64 = image_to_base64(logo_path)
    title_page_html = f"""
    <div style='text-align:center'>
        <h1 style='font-size:36px; font-weight:bold;'>{project_name}</h1>
        <img src='data:image/jpeg;base64,{logo_base64}' width='150'>
        <table style='width:50%; margin-top:20px; margin-left:auto; margin-right:auto; border-collapse: collapse; border: 1px solid black;'>
            <thead>
                <tr>
                    <th style='border: 1px solid black; padding: 8px;'>Version</th>
                    <th style='border: 1px solid black; padding: 8px;'>Date</th>
                </tr>
            </thead>
            <tbody>
                <tr>
                    <td style='border: 1px solid black; padding: 8px;'>1.0</td>
                    <td style='border: 1px solid black; padding: 8px;'>dd/mm/yyyy</td>
                </tr>
            </tbody>
        </table>
    </div>
    <div style='page-break-after: always;'></div>  <!-- Page break -->
    """
    return title_page_html + html_text




def convert_markdown_to_pdf_and_word(markdown_file_path, pdf_file_path, word_file_path):
    os.makedirs(os.path.dirname(pdf_file_path), exist_ok=True)
    os.makedirs(os.path.dirname(word_file_path), exist_ok=True)

    with open(markdown_file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    project_name = extract_project_name_from_markdown(md_text)

    html_text = markdown.markdown(md_text)

    doc = Document()

    if project_name:
        add_title_page_to_word(doc, project_name, "Prompt2RFP/Company Logo/55/logo.jpg")
        add_footer_to_word(doc, "FiftyFive Technologies Pvt Ltd")

    soup = BeautifulSoup(html_text, 'html.parser')

    for element in soup:
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            p = doc.add_paragraph()
            run = p.add_run(element.text)
            run.bold = True
            run.font.size = Pt(16 if element.name in ['h1', 'h2'] else 14)
        elif element.name == 'p':
            p = doc.add_paragraph(element.text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif element.name == 'ul':
            for item in element.find_all('li'):
                p = doc.add_paragraph()
                p.add_run('• ')
                p.add_run(item.text)
        elif element.name == 'ol':
            for idx, item in enumerate(element.find_all('li'), 1):
                p = doc.add_paragraph()
                p.add_run(f'{idx}. {item.text}')

    doc.save(word_file_path)

    markdown_text = add_line_breaks_between_list_items(md_text)
    html_text = markdown.markdown(markdown_text)

    if project_name:
        wrapped_html = add_title_page_to_pdf(html_text, project_name, "Prompt2RFP/Company Logo/55/logo.jpg")
    else:
        wrapped_html = html_text

    pdfkit.from_string(wrapped_html, pdf_file_path, {
        'margin-top': '20mm',
        'margin-bottom': '15mm',
        'margin-right': '10mm',
        'margin-left': '10mm',
        'footer-center': '[page] of [topage]',
        'footer-left': 'FiftyFive Technologies Pvt Ltd',
    })
