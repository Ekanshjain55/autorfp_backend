from docx import Document

# Step 1: Read the Markdown file
def read_markdown_file(filename):
    with open(filename, "r") as f:
        return f.read()

# Step 2: Parse the Markdown file to extract sections and their content
def parse_markdown_content(markdown_content):
    sections = markdown_content.split("\n# ")
    section_dict = {}
    for section in sections:
        lines = section.split("\n")
        section_title = lines[0].replace("# ", "")
        section_content = "\n".join(lines[1:])
        section_dict[section_title] = section_content
    return section_dict

# Step 3: Read the Word template and Step 4: Populate the Word template
def populate_word_template(template_path, section_dict):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in section_dict.items():
            if key in paragraph.text:
                paragraph.clear()
                paragraph.add_run(value)
    return doc

# Step 5: Save the populated Word Document
def save_populated_word_document(doc, save_path):
    doc.save(save_path)

if __name__ == "__main__":
    # Define file paths
    markdown_file_path = "/home/jayantnehra/RFP_doc/Generated Proposal/generated_rfp.md"
    word_template_path = "/home/jayantnehra/RFP_doc/Project_Proposals/Project_Proposal_Template.docx"
    populated_word_save_path = "/home/jayantnehra/RFP_doc/Generated Proposal/generated_rfp_docparsed.docx"
    
    # Read the Markdown file
    markdown_content = read_markdown_file(markdown_file_path)
    
    # Parse the Markdown content
    section_dict = parse_markdown_content(markdown_content)
    
    # Populate the Word template
    doc = populate_word_template(word_template_path, section_dict)
    
    # Save the populated Word document
    save_populated_word_document(doc, populated_word_save_path)
